import sys
from pathlib import Path

# Add project root directory to sys.path
PROJECT_ROOT = Path(__file__).resolve().parents[1]  # Adjust parent index if needed
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

import os
import json
import time
import uuid
import anthropic
from typing import Any
from dotenv import load_dotenv
from agents.tools import execute_tool, RESEARCHER_TOOLS, VALIDATOR_TOOLS
from orchestrator.state import AgentState

load_dotenv()

client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))


# ── Helper: agent loop with tool use ─────────────────────────────────────────

def run_agent_with_tools(
    system_prompt: str,
    user_message: str,
    tools: list,
    model: str = "claude-sonnet-4-6",
    max_iterations: int = 6
) -> tuple[str, list]:
    """
    Runs the tool-use loop for a single agent node.
    Returns (final_text_response, list_of_tool_calls_made).
    """
    messages = [{"role": "user", "content": user_message}]
    tool_calls_log = []
    iterations = 0

    while iterations < max_iterations:
        iterations += 1
        response = client.messages.create(
            model=model,
            max_tokens=2048,
            system=system_prompt,
            tools=tools,
            messages=messages
        )

        if response.stop_reason == "end_turn":
            # Agent is done — extract text response
            final_text = next(
                (block.text for block in response.content if hasattr(block, "text")),
                ""
            )
            return final_text, tool_calls_log

        if response.stop_reason == "tool_use":
            # Execute all tool calls in this response
            tool_results = []
            for block in response.content:
                if block.type == "tool_use":
                    tool_name = block.name
                    tool_input = block.input
                    tool_use_id = block.id

                    result_str = execute_tool(tool_name, tool_input)
                    tool_calls_log.append({
                        "tool": tool_name,
                        "input": tool_input,
                        "result_preview": result_str[:200]
                    })

                    tool_results.append({
                        "type": "tool_result",
                        "tool_use_id": tool_use_id,
                        "content": result_str
                    })

            # Append assistant message + tool results to history
            messages.append({"role": "assistant", "content": response.content})
            messages.append({"role": "user", "content": tool_results})

        else:
            # Unexpected stop reason
            break

    # Max iterations reached — return what we have
    final_text = next(
        (block.text for block in response.content if hasattr(block, "text")),
        "Agent reached maximum iterations without completing."
    )
    return final_text, tool_calls_log


# ── Node 1: Classifier ────────────────────────────────────────────────────────

def classify_node(state: AgentState) -> dict:
    """
    Classifies the defect by category, severity, and whether it's a known issue.
    Produces search keywords for the researcher.
    No tools needed — pure reasoning.
    """
    defect = state["defect_input"]

    system = """You are an automotive defect classification expert at BMW Group.
Your job is to analyse an incoming defect report and classify it so the research
and documentation team can efficiently find the relevant information.

You MUST respond with a valid JSON object and nothing else. No explanation, no markdown.
Schema:
{
  "defect_category": string,     // e.g. "engine", "transmission", "electrical", "brake", "fuel_system"
  "severity": string,            // "low" | "medium" | "high" | "safety_critical"
  "is_known_issue": boolean,     // true if this matches a common pattern you recognise
  "search_keywords": [string],   // 4-6 specific technical keywords for documentation search
  "reasoning": string            // one sentence explaining your classification
}"""

    user_message = f"""Classify this incoming defect report:

Vehicle: {defect['vehicle_model']} ({defect['year']})
Mileage: {defect['mileage_km']:,} km
Error codes: {', '.join(defect['error_codes']) if defect['error_codes'] else 'None reported'}
Symptom: {defect['symptom_description']}"""

    response = client.messages.create(
        model="claude-haiku-4-5",   # Fast, cheap for classification
        max_tokens=512,
        system=system,
        messages=[{"role": "user", "content": user_message}]
    )

    raw = response.content[0].text.strip()

    # Strip markdown fences if model adds them
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]

    try:
        classification = json.loads(raw)
    except json.JSONDecodeError:
        # Fallback if model doesn't return clean JSON
        classification = {
            "defect_category": "unknown",
            "severity": "medium",
            "is_known_issue": False,
            "search_keywords": defect["error_codes"] + [defect["symptom_description"][:50]],
            "reasoning": "Classification failed — defaulting to medium severity."
        }

    return {"classification": classification}


# ── Node 2: Documentation Researcher ─────────────────────────────────────────

def research_node(state: AgentState) -> dict:
    """
    Uses RAG to find relevant repair procedures, TSBs, and known fixes.
    Tools: search_vector_store, fetch_document_section.
    """
    defect = state["defect_input"]
    classification = state["classification"]
    keywords = ", ".join(classification["search_keywords"])
    feedback = state.get("human_feedback") or ""

    system = """You are an automotive documentation researcher at BMW Group.
Your role is to find the most relevant technical information for a vehicle defect
from the internal knowledge base.

Use the search_vector_store tool to retrieve relevant documentation.
If a search result looks highly relevant but you need more context, use
fetch_document_section to read the full page.

After your research, respond with a JSON object (no markdown, no explanation):
{
  "relevant_documents": [{"title": string, "source": string, "key_finding": string}],
  "root_cause_hypothesis": string,
  "recommended_procedure": string,
  "required_parts": [string],
  "estimated_repair_time_hours": number,
  "sources": [string]
}"""

    feedback_section = f"\n\nIMPORTANT — Human reviewer feedback on previous attempt:\n{feedback}\nAddress these concerns in your research." if feedback else ""

    user_message = f"""Research this vehicle defect:
                    Vehicle: {defect['vehicle_model']} ({defect['year']}), {defect['mileage_km']:,} km
                    Error codes: {', '.join(defect['error_codes']) if defect['error_codes'] else 'None'}
                    Symptom: {defect['symptom_description']}
                    Classification: {classification['defect_category']} — {classification['severity']} severity
                    Search keywords: {keywords}{feedback_section}

                    Search for TSBs, manuals, repair procedures, and known issues related to this defect.
                    Be thorough — search with multiple queries covering different angles of the problem. 
                  """

    raw_response, tool_calls = run_agent_with_tools(
        system_prompt=system,
        user_message=user_message,
        tools=RESEARCHER_TOOLS,
        model="claude-sonnet-4-6"
    )

    text = raw_response.strip()
    if text.startswith("```"):
        text = text.split("```")[1]
        if text.startswith("json"):
            text = text[4:]

    try:
        research = json.loads(text)
    except json.JSONDecodeError:
        research = {
            "relevant_documents": [],
            "root_cause_hypothesis": raw_response[:500],
            "recommended_procedure": "See raw response — JSON parsing failed.",
            "required_parts": [],
            "estimated_repair_time_hours": 0.0,
            "sources": []
        }

    return {
        "research": research,
        "iteration_count": state.get("iteration_count", 0) + 1
    }


# ── Node 3: Standards Validator ───────────────────────────────────────────────

def validate_node(state: AgentState) -> dict:
    """
    Checks the proposed repair procedure against automotive standards.
    Tools: search_standards_db, web_search.
    """
    research = state["research"]
    classification = state["classification"]

    system = """You are a quality and compliance specialist at BMW Group.
Your role is to verify that proposed repair procedures comply with relevant
automotive standards (IATF 16949, ISO 9001, BMW group standards).

Use search_standards_db to find relevant clauses and requirements.
Use web_search only if local results are insufficient or you need to check
for very recent regulatory updates.

Respond with a JSON object (no markdown, no explanation):
{
  "complies_with_standards": boolean,
  "applicable_standards": [string],
  "compliance_notes": string,
  "warnings": [string],
  "requires_escalation": boolean
}"""

    user_message = f"""Validate this proposed repair procedure against automotive standards:

Defect category: {classification['defect_category']}
Severity: {classification['severity']}
Root cause hypothesis: {research['root_cause_hypothesis']}
Proposed procedure: {research['recommended_procedure']}
Required parts: {', '.join(research['required_parts']) if research['required_parts'] else 'None specified'}

Check compliance with IATF 16949 and any other applicable standards.
Flag any safety or regulatory concerns."""

    raw_response, tool_calls = run_agent_with_tools(
        system_prompt=system,
        user_message=user_message,
        tools=VALIDATOR_TOOLS,
        model="claude-sonnet-4-6"
    )

    text = raw_response.strip()
    if text.startswith("```"):
        text = text.split("```")[1]
        if text.startswith("json"):
            text = text[4:]

    try:
        validation = json.loads(text)
    except json.JSONDecodeError:
        validation = {
            "complies_with_standards": False,
            "applicable_standards": [],
            "compliance_notes": "Validation parsing failed — manual review required.",
            "warnings": ["JSON parse error in validator output"],
            "requires_escalation": True
        }

    return {"validation": validation}


# ── Node 4: Report Writer ─────────────────────────────────────────────────────

def write_report_node(state: AgentState) -> dict:
    """
    Synthesises all agent outputs into a structured final report.
    No tools — pure generation from accumulated context.
    """
    defect = state["defect_input"]
    classification = state["classification"]
    research = state["research"]
    validation = state["validation"]

    severity_to_priority = {
        "low": "routine",
        "medium": "routine",
        "high": "urgent",
        "safety_critical": "immediate"
    }

    system = """You are a senior technical writer at BMW Group.
Synthesise the research and validation findings into a clean, structured defect report.
Respond ONLY with a valid JSON object matching the schema exactly."""

    user_message = f"""Create a defect report from these findings:

INPUT:
Vehicle: {defect['vehicle_model']} ({defect['year']}), {defect['mileage_km']:,} km
Error codes: {defect['error_codes']}
Symptom: {defect['symptom_description']}

CLASSIFICATION:
{json.dumps(classification, indent=2)}

RESEARCH FINDINGS:
{json.dumps(research, indent=2)}

STANDARDS VALIDATION:
{json.dumps(validation, indent=2)}

Return a JSON object with these exact fields:
{{
  "report_id": "{state['report_id']}",
  "vehicle_info": {{"model": string, "year": int, "mileage_km": int, "error_codes": [string]}},
  "defect_category": string,
  "severity": string,
  "root_cause": string,
  "recommended_action": string,
  "required_parts": [string],
  "estimated_hours": number,
  "standards_compliance": string,
  "warnings": [string],
  "sources": [string],
  "priority": string
}}"""

    response = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=1024,
        system=system,
        messages=[{"role": "user", "content": user_message}]
    )

    raw = response.content[0].text.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]

    try:
        report = json.loads(raw)
    except json.JSONDecodeError:
        report = {
            "report_id": state["report_id"],
            "vehicle_info": defect,
            "defect_category": classification.get("defect_category", "unknown"),
            "severity": classification.get("severity", "medium"),
            "root_cause": research.get("root_cause_hypothesis", "Unknown"),
            "recommended_action": research.get("recommended_procedure", "Manual review required"),
            "required_parts": research.get("required_parts", []),
            "estimated_hours": research.get("estimated_repair_time_hours", 0),
            "standards_compliance": validation.get("compliance_notes", "Unknown"),
            "warnings": validation.get("warnings", []),
            "sources": research.get("sources", []),
            "priority": severity_to_priority.get(classification.get("severity", "medium"), "routine")
        }

    return {"report": report}


# ── Node 5: Human Checkpoint ──────────────────────────────────────────────────

def human_checkpoint_node(state: AgentState) -> dict:
    """
    Pauses the graph for human review.
    LangGraph's interrupt() suspends execution here.
    The graph resumes when the FastAPI endpoint sends a Command(resume=...).
    """
    from langgraph.types import interrupt

    decision = interrupt({
        "message": "Please review the generated report and approve or request revision.",
        "report": state["report"],
        "iteration": state.get("iteration_count", 1)
    })

    return {
        "human_approved": decision.get("approved", False),
        "human_feedback": decision.get("feedback", "")
    }


# ── Node 6: Save Report ───────────────────────────────────────────────────────

def save_report_node(state: AgentState) -> dict:
    """
    Saves the approved report to PostgreSQL and generates a Word document.
    """
    import json
    from docx import Document
    import sqlite3
    from pathlib import Path

    report = state["report"]

    # Save to sqlLite
    DATABASE_PATH = r"C:\Users\Bakwowi Junior\Documents\My-Portfolio\Automotive Process Intelligence Agent\data\sqlLite_db\apia_db.db"

    with sqlite3.connect(DATABASE_PATH) as conn:
        cursor = conn.cursor()
        cursor.execute("""
            UPDATE defect_reports
            SET report_data = :report_data,
                status = 'approved',
                updated_at = CURRENT_TIMESTAMP
            WHERE report_id = :report_id
        """, {
            "report_id": report["report_id"],
            "report_data": json.dumps(report)
        })
        conn.commit()

    # Generate Word document
    doc = Document()
    doc.add_heading("BMW Group — Defect Report", 0)
    doc.add_heading(f"Report ID: {report['report_id']}", level=1)

    # Vehicle info
    doc.add_heading("Vehicle Information", level=2)
    vi = report["vehicle_info"]
    doc.add_paragraph(f"Model: {vi.get('model', vi.get('vehicle_model', 'N/A'))}")
    doc.add_paragraph(f"Year: {vi.get('year', 'N/A')}")
    doc.add_paragraph(f"Mileage: {vi.get('mileage_km', vi.get('mileage', 'N/A')):,} km")
    doc.add_paragraph(f"Error Codes: {', '.join(vi.get('error_codes', []))}")

    # Classification
    doc.add_heading("Classification", level=2)
    doc.add_paragraph(f"Category: {report['defect_category']}")
    doc.add_paragraph(f"Severity: {report['severity'].upper()}")
    doc.add_paragraph(f"Priority: {report['priority'].upper()}")

    # Findings
    doc.add_heading("Root Cause", level=2)
    doc.add_paragraph(report["root_cause"])

    doc.add_heading("Recommended Action", level=2)
    doc.add_paragraph(report["recommended_action"])

    # Parts
    if report["required_parts"]:
        doc.add_heading("Required Parts", level=2)
        for part in report["required_parts"]:
            doc.add_paragraph(f"• {part}")

    doc.add_paragraph(f"Estimated Repair Time: {report['estimated_hours']} hours")

    # Compliance
    doc.add_heading("Standards Compliance", level=2)
    doc.add_paragraph(report["standards_compliance"])
    if report["warnings"]:
        doc.add_heading("Warnings", level=3)
        for warning in report["warnings"]:
            doc.add_paragraph(f"⚠ {warning}")

    # Sources
    if report["sources"]:
        doc.add_heading("Sources", level=2)
        for source in report["sources"]:
            doc.add_paragraph(f"• {source}")

    # Save file
    output_dir = Path("output/reports")
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"report_{report['report_id']}.docx"
    doc.save(str(output_path))

    return {"error_log": state.get("error_log", []) + [f"Report saved to {output_path}"]}